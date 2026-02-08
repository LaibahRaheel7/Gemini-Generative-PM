"""
Document Parser - Extracts text from PDF and DOCX files for AI analysis.
"""
import os
from typing import Optional
from pathlib import Path


class DocumentParser:
    """
    Handles text extraction from various document formats.
    Supports PDF and DOCX files.
    """
    
    def __init__(self):
        """Initialize the document parser."""
        pass
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """
        Extract text content from a PDF file.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Extracted text content
            
        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If file is not a valid PDF
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"PDF file not found: {file_path}")
        
        try:
            from PyPDF2 import PdfReader
            
            reader = PdfReader(file_path)
            text_content = []
            
            # Extract text from each page
            for page_num, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(f"--- Page {page_num + 1} ---\n{page_text}")
                except Exception as e:
                    print(f"Warning: Could not extract text from page {page_num + 1}: {e}")
            
            if not text_content:
                return "Warning: No text could be extracted from the PDF. It may contain only images."
            
            return "\n\n".join(text_content)
            
        except ImportError:
            raise ImportError("PyPDF2 is not installed. Run: pip install PyPDF2==3.0.1")
        except Exception as e:
            raise ValueError(f"Failed to parse PDF: {str(e)}")
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """
        Extract text content from a DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Extracted text content
            
        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If file is not a valid DOCX
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"DOCX file not found: {file_path}")
        
        try:
            from docx import Document
            
            doc = Document(file_path)
            text_content = []
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_content.append(row_text)
            
            if not text_content:
                return "Warning: No text could be extracted from the DOCX file."
            
            return "\n\n".join(text_content)
            
        except ImportError:
            raise ImportError("python-docx is not installed. Run: pip install python-docx==1.1.0")
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX: {str(e)}")
    
    def extract_text(self, file_path: str) -> str:
        """
        Auto-detect file type and extract text.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Extracted text content
            
        Raises:
            ValueError: If file format is not supported
        """
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            if file_ext == '.doc':
                raise ValueError(
                    "Legacy .doc format is not supported. "
                    "Please convert to .docx format first."
                )
            return self.extract_text_from_docx(file_path)
        else:
            raise ValueError(
                f"Unsupported file format: {file_ext}. "
                f"Supported formats: .pdf, .docx"
            )
    
    def extract_text_from_bytes(self, file_bytes: bytes, filename: str) -> str:
        """
        Extract text from file bytes (useful for uploaded files).
        
        Args:
            file_bytes: Bytes content of the file
            filename: Original filename (used to determine format)
            
        Returns:
            Extracted text content
        """
        import tempfile
        
        # Write bytes to a temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix=Path(filename).suffix) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name
        
        try:
            # Extract text from the temporary file
            return self.extract_text(tmp_path)
        finally:
            # Clean up temporary file
            try:
                os.unlink(tmp_path)
            except Exception:
                pass
    
    def validate_file_size(self, file_path: str, max_size_mb: int = 50) -> bool:
        """
        Validate that file size is within acceptable limits.
        
        Args:
            file_path: Path to the file
            max_size_mb: Maximum allowed size in megabytes
            
        Returns:
            True if file size is acceptable
            
        Raises:
            ValueError: If file is too large
        """
        file_size = os.path.getsize(file_path)
        file_size_mb = file_size / (1024 * 1024)
        
        if file_size_mb > max_size_mb:
            raise ValueError(
                f"File is too large: {file_size_mb:.1f}MB. "
                f"Maximum allowed size is {max_size_mb}MB."
            )
        
        return True
    
    def get_document_info(self, file_path: str) -> dict:
        """
        Get metadata about a document.
        
        Args:
            file_path: Path to the document
            
        Returns:
            Dictionary with document metadata
        """
        file_ext = Path(file_path).suffix.lower()
        file_size = os.path.getsize(file_path)
        
        info = {
            "filename": Path(file_path).name,
            "extension": file_ext,
            "size_bytes": file_size,
            "size_mb": file_size / (1024 * 1024),
            "supported": file_ext in ['.pdf', '.docx']
        }
        
        # Get page count for PDFs
        if file_ext == '.pdf':
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(file_path)
                info["page_count"] = len(reader.pages)
            except Exception:
                info["page_count"] = None
        
        # Get paragraph count for DOCX
        elif file_ext == '.docx':
            try:
                from docx import Document
                doc = Document(file_path)
                info["paragraph_count"] = len(doc.paragraphs)
                info["table_count"] = len(doc.tables)
            except Exception:
                info["paragraph_count"] = None
                info["table_count"] = None
        
        return info


# Singleton instance
_document_parser_instance = None


def get_document_parser() -> DocumentParser:
    """
    Get the singleton DocumentParser instance.
    
    Returns:
        DocumentParser instance
    """
    global _document_parser_instance
    if _document_parser_instance is None:
        _document_parser_instance = DocumentParser()
    return _document_parser_instance
